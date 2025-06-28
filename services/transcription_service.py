from clients.gmail_client import GmailClient
from clients.calendar_client import CalendarClient
from clients.github_client import GitHubClient
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
import re
from yaml import safe_dump
import base64
import json


class TranscriptionService:
    def __init__(self):
        self.gmail = GmailClient()
        self.calendar = CalendarClient()
        self.github = GitHubClient()

    def run(self) -> list:
        agenda_map = self.calendar.get_events_map()
        return self.lire_emails_fathom(agenda_map)

    def sync_transcriptions(self, limit=None):
        transcriptions = self.lire_emails_fathom(self.calendar.get_events_map(), limit=limit)
        index = []
        updates = {}

        for r in transcriptions:
            filename = f"MH/transcriptions/{r['date'].replace('/', '-')}_{r['titre'].replace(' ', '_')}.md"
            content = self._generate_markdown(r)
            updates[filename] = content

            index.append({
                "id": filename.split("/")[-1].replace(".md", ""),
                "titre": r["titre"],
                "date": r["date"],
                "heure": r["heure"],
                "fichier": filename,
                "status": "en_cours",
                "tags": ["transcription", "réunion"]
            })

        # Lire index.json existant et le mettre à jour
        index_path = "MH/index.json"
        try:
            existing = self.github.read_json(index_path)
            existing_ids = {item["id"] for item in existing}
            updated = existing + [item for item in index if item["id"] not in existing_ids]
        except Exception:
            updated = index

        updates[index_path] = self._generate_index_json(updated)

        # Commit unique de tous les fichiers
        self.github.upload_files(updates, commit_message="Ajout de transcriptions Fathom + mise à jour index")
        return [i["fichier"] for i in index]

    def lire_emails_fathom(self, agenda_map, limit=10):
        emails = self.gmail.rechercher("from:no-reply@fathom.video", max_results=limit)
        reunions = []
        for email in emails:
            full_msg = self.gmail.get_message(email["id"])
            if not full_msg:
                continue

            payload = full_msg.get("payload", {})
            texte = self._decode_message(payload)
            texte = self._nettoyer_html(texte)

            titre = self._extraire_titre_email(texte)
            event = agenda_map.get(titre, {})
            reunions.append({
                "titre": titre.replace("TR :", "").replace("TR:", "").strip(),
                "date": event.get("date", "(Date inconnue)"),
                "heure": event.get("heure", "(Horaire inconnu)"),
                "conclusions": self._extraire_section(texte, "Principales conclusions", "Sujets abordés"),
                "sujets": self._extraire_section(texte, "Sujets abordés", "Prochaines étapes"),
                "prochaines étapes": self._extraire_section(texte, "Prochaines étapes"),
                "texte_complet": texte
            })
        return reunions

    def _decode_message(self, payload):
        if 'parts' in payload:
            for part in payload['parts']:
                if part.get('mimeType') in ['text/plain', 'text/html']:
                    data = part['body'].get('data', '')
                    if data:
                        decoded = base64.urlsafe_b64decode(data).decode('utf-8')
                        return decoded
                elif 'parts' in part:
                    return self._decode_message(part)
        elif payload.get('mimeType') in ['text/plain', 'text/html']:
            data = payload['body'].get('data', '')
            if data:
                return base64.urlsafe_b64decode(data).decode('utf-8')
        return ""

    def _nettoyer_html(self, html):
        return BeautifulSoup(html, "html.parser").get_text(separator="\n", strip=True)

    def _extraire_titre_email(self, texte):
        match = re.search(r"Meeting with\s+(.*?)\s*Objectif de la réunion", texte, re.DOTALL | re.IGNORECASE)
        if match:
            bloc = match.group(1).strip().splitlines()
            bloc = [l.strip() for l in bloc if l.strip()]
            if len(bloc) >= 2:
                titre = bloc[1]
                return titre.strip()
        return "(Titre non trouvé)"

    def _extraire_section(self, texte, debut, fin=None):
        pattern = rf"(?:✅\s*)?{re.escape(debut)}\s*:?(.*?)(?:{re.escape(fin)}\s*:?|$)" if fin \
                  else rf"(?:✅\s*)?{re.escape(debut)}\s*:?(.*?)(?:Meeting with|$)"
        match = re.search(pattern, texte, re.DOTALL | re.IGNORECASE)
        return match.group(1).strip() if match and match.group(1).strip() else "(Non renseigné)"

    def _generate_markdown(self, r) -> str:
        metadata = {
            "titre": r["titre"],
            "date": r["date"],
            "heure": r["heure"],
            "status": "en_cours",
            "tags": ["transcription", "réunion"]
        }

        frontmatter = safe_dump(metadata, allow_unicode=True, sort_keys=False)
        body = f"## 🧠 Principales conclusions\n{r['conclusions']}\n\n" \
               f"## 📌 Sujets abordés\n{r['sujets']}\n\n" \
               f"## 🧭 Prochaines étapes\n{r.get('prochaines étapes', '(Non renseigné)')}"
        return f"---\n{frontmatter}---\n\n{body.strip()}"

    def _generate_index_json(self, data: list) -> str:
        return json.dumps(data, indent=2, ensure_ascii=False)

    def generer_docx_reunion(self, r):
        doc = Document()
        doc.add_heading(f"Réunion : {r['titre']} – {r['date']} ({r['heure']})", level=1)

        def ajouter_section(titre, contenu):
            doc.add_heading(titre, level=2)
            blocs = contenu.strip().split("\n\n")
            for bloc in blocs:
                lignes = [l.strip() for l in bloc.strip().splitlines() if l.strip()]
                if len(lignes) == 1:
                    doc.add_paragraph(lignes[0], style="Normal")
                else:
                    doc.add_paragraph(lignes[0], style='Intense Quote')
                    for ligne in lignes[1:]:
                        if ligne.startswith("- "):
                            doc.add_paragraph(ligne[2:], style="List Bullet")
                        else:
                            doc.add_paragraph(ligne, style="Body Text")

        ajouter_section("Principales conclusions", r["conclusions"])
        ajouter_section("Sujets abordés", r["sujets"])
        if r.get("prochaines étapes"):
            ajouter_section("Prochaines étapes", r["prochaines étapes"])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
